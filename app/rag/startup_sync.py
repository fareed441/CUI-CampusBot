"""
CUI CampusBot - Startup Sync Module
Syncs MongoDB documents to ChromaDB on server startup

This module ensures MongoDB is the SINGLE SOURCE OF TRUTH for documents.
ChromaDB is rebuilt/synced from MongoDB on each startup.
"""

import logging
from datetime import datetime
from typing import Optional, Dict, List
from bson import ObjectId

from app.database import get_database
from app.models import DocumentStatus
from app.config import EMBEDDING_MODEL, CHUNK_SIZE, CHUNK_OVERLAP, CHROMA_PERSIST_DIRECTORY

logger = logging.getLogger(__name__)


class MongoDBToChromaSync:
    """
    Handles synchronization of documents from MongoDB/GridFS to ChromaDB.
    """
    
    def __init__(self):
        self.db = None
        self.vectorstore = None
        self.embeddings = None
        
    def initialize_embeddings(self):
        """Initialize multilingual embeddings model"""
        from langchain_huggingface import HuggingFaceEmbeddings
        
        logger.info(f"Initializing embeddings model: {EMBEDDING_MODEL}")
        
        self.embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True}
        )
        
        logger.info("[OK] Multilingual embeddings model initialized")
        return self.embeddings
    
    def initialize_vectorstore(self):
        """Initialize or connect to ChromaDB"""
        from langchain_chroma import Chroma
        
        if self.embeddings is None:
            self.initialize_embeddings()
        
        self.vectorstore = Chroma(
            collection_name="cui_knowledge_base",
            embedding_function=self.embeddings,
            persist_directory=CHROMA_PERSIST_DIRECTORY
        )
        
        logger.info(f"[OK] ChromaDB initialized at: {CHROMA_PERSIST_DIRECTORY}")
        return self.vectorstore
    
    def get_mongodb_documents(self) -> List[dict]:
        """Get all documents from MongoDB"""
        self.db = get_database()
        documents = list(self.db.documents.find())
        logger.info(f"Found {len(documents)} documents in MongoDB")
        return documents
    
    def get_chromadb_document_ids(self) -> set:
        """Get set of document IDs already in ChromaDB"""
        if self.vectorstore is None:
            self.initialize_vectorstore()
        
        try:
            results = self.vectorstore._collection.get(include=["metadatas"])
            doc_ids = set()
            
            if results and results.get("metadatas"):
                for metadata in results["metadatas"]:
                    if metadata and "document_id" in metadata:
                        doc_ids.add(metadata["document_id"])
            
            return doc_ids
        except Exception as e:
            logger.warning(f"Could not fetch ChromaDB document IDs: {e}")
            return set()
    
    def load_file_from_gridfs(self, gridfs_file_id: str) -> Optional[bytes]:
        """Load file content from MongoDB GridFS"""
        try:
            file_data = self.db.get_file(gridfs_file_id)
            return file_data
        except Exception as e:
            logger.error(f"Failed to load file from GridFS: {e}")
            return None
    
    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text from document based on file type"""
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        try:
            text = ""
            
            if file_type == "pdf":
                try:
                    import pdfplumber
                    with pdfplumber.open(tmp_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                except ImportError:
                    from pypdf import PdfReader
                    reader = PdfReader(tmp_path)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            
            elif file_type == "docx":
                from docx import Document
                doc = Document(tmp_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                        
            elif file_type == "txt":
                with open(tmp_path, "r", encoding="utf-8") as f:
                    text = f.read()
                    
            elif file_type == "json":
                import json
                with open(tmp_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                text = self._json_to_text(data)
            
            return text.strip()
            
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    def _json_to_text(self, data, prefix="") -> str:
        """Convert JSON data to readable text"""
        text = ""
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text += self._json_to_text(value, f"{prefix}{key}: ")
                else:
                    text += f"{prefix}{key}: {value}\n"
        elif isinstance(data, list):
            for i, item in enumerate(data):
                text += self._json_to_text(item, f"{prefix}[{i}] ")
        else:
            text += f"{prefix}{data}\n"
        
        return text
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks for embedding"""
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        return text_splitter.split_text(text)
    
    def add_document_to_chromadb(self, document_id: str, chunks: List[str], metadata: dict) -> int:
        """Add document chunks to ChromaDB"""
        if self.vectorstore is None:
            self.initialize_vectorstore()
        
        ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]
        
        metadatas = [
            {
                "document_id": document_id,
                "chunk_index": i,
                "title": metadata.get("title", ""),
                "file_name": metadata.get("file_name", ""),
                "file_type": metadata.get("file_type", "")
            }
            for i in range(len(chunks))
        ]
        
        self.vectorstore.add_texts(texts=chunks, ids=ids, metadatas=metadatas)
        return len(chunks)
    
    def process_document(self, doc: dict) -> bool:
        """Process a single document from MongoDB to ChromaDB"""
        doc_id = str(doc["_id"])
        
        try:
            self.db.documents.update_one(
                {"_id": doc["_id"]},
                {"$set": {"status": DocumentStatus.PROCESSING.value}}
            )
            
            file_content = self.load_file_from_gridfs(doc["gridfs_file_id"])
            if not file_content:
                raise ValueError(f"Could not load file from GridFS")
            
            text = self.extract_text(file_content, doc["file_type"])
            if not text:
                raise ValueError("No text content extracted")
            
            chunks = self.chunk_text(text)
            if not chunks:
                raise ValueError("No chunks generated")
            
            chunk_count = self.add_document_to_chromadb(doc_id, chunks, doc)
            
            self.db.documents.update_one(
                {"_id": doc["_id"]},
                {
                    "$set": {
                        "status": DocumentStatus.PROCESSED.value,
                        "chunk_count": chunk_count,
                        "embedding_model": EMBEDDING_MODEL,
                        "processed_at": datetime.utcnow()
                    }
                }
            )
            
            logger.info(f"[OK] Processed: {doc['file_name']} ({chunk_count} chunks)")
            return True
            
        except Exception as e:
            logger.error(f"Failed to process {doc.get('file_name')}: {e}")
            self.db.documents.update_one(
                {"_id": doc["_id"]},
                {"$set": {"status": DocumentStatus.FAILED.value, "error_message": str(e)}}
            )
            return False
    
    def sync_all(self, force_rebuild: bool = False) -> Dict:
        """Sync all MongoDB documents to ChromaDB"""
        logger.info("=" * 60)
        logger.info("STARTING MONGODB TO CHROMADB SYNC")
        logger.info("=" * 60)
        
        results = {
            "total_documents": 0,
            "already_synced": 0,
            "newly_processed": 0,
            "failed": 0,
            "errors": []
        }
        
        self.initialize_embeddings()
        self.initialize_vectorstore()
        
        documents = self.get_mongodb_documents()
        results["total_documents"] = len(documents)
        
        if not documents:
            logger.info("No documents found in MongoDB. Sync complete.")
            return results
        
        existing_ids = self.get_chromadb_document_ids() if not force_rebuild else set()
        
        for doc in documents:
            doc_id = str(doc["_id"])
            
            if not force_rebuild and doc_id in existing_ids and doc.get("status") == DocumentStatus.PROCESSED.value:
                logger.info(f"[SKIP] Already synced: {doc['file_name']}")
                results["already_synced"] += 1
                continue
            
            success = self.process_document(doc)
            
            if success:
                results["newly_processed"] += 1
            else:
                results["failed"] += 1
                results["errors"].append(doc.get("file_name", "unknown"))
        
        logger.info("=" * 60)
        logger.info("MONGODB TO CHROMADB SYNC COMPLETE")
        logger.info(f"Total: {results['total_documents']}, Synced: {results['already_synced']}, New: {results['newly_processed']}, Failed: {results['failed']}")
        logger.info("=" * 60)
        
        return results
    
    def get_vectorstore(self):
        """Get the initialized vectorstore"""
        if self.vectorstore is None:
            self.initialize_vectorstore()
        return self.vectorstore
    
    def get_embeddings(self):
        """Get the initialized embeddings model"""
        if self.embeddings is None:
            self.initialize_embeddings()
        return self.embeddings


_sync_instance: Optional[MongoDBToChromaSync] = None


def get_sync_instance() -> MongoDBToChromaSync:
    """Get or create the global sync instance"""
    global _sync_instance
    if _sync_instance is None:
        _sync_instance = MongoDBToChromaSync()
    return _sync_instance


def sync_mongodb_to_chromadb(force_rebuild: bool = False) -> Dict:
    """Main function to sync MongoDB documents to ChromaDB"""
    sync = get_sync_instance()
    return sync.sync_all(force_rebuild)


def get_vectorstore():
    """Get the ChromaDB vectorstore for RAG queries"""
    sync = get_sync_instance()
    return sync.get_vectorstore()


def get_embeddings():
    """Get the embeddings model for queries"""
    sync = get_sync_instance()
    return sync.get_embeddings()
