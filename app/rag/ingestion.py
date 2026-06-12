"""
CUI CampusBot - RAG Ingestion Module
Document processing and ChromaDB embedding pipeline
"""

from datetime import datetime
from typing import Optional
import logging
import tempfile
import os
from bson import ObjectId

from app.database import get_database
from app.config import EMBEDDING_MODEL, CHUNK_SIZE, CHUNK_OVERLAP, CHROMA_PERSIST_DIRECTORY
from app.models import DocumentStatus

logger = logging.getLogger(__name__)

# ===========================================
# Document Processing
# ===========================================

def process_document(document_id: str, file_content: bytes, file_type: str) -> bool:
    """
    Process a document through the RAG pipeline
    
    1. Extract text from document
    2. Split into chunks
    3. Generate embeddings
    4. Store in ChromaDB
    
    Args:
        document_id: MongoDB document ID
        file_content: Raw file bytes
        file_type: File extension (pdf, docx, txt, json)
    
    Returns:
        True if successful
    """
    db = get_database()
    
    try:
        # Update status to processing
        db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {"$set": {"status": DocumentStatus.PROCESSING.value}}
        )
        
        logger.info(f"Processing document: {document_id}")
        
        # Save to temp file for processing
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        try:
            # Extract text based on file type
            text_content = extract_text(tmp_path, file_type)
            
            if not text_content:
                raise ValueError("No text content extracted from document")
            
            # Split into chunks
            chunks = split_into_chunks(text_content)
            
            if not chunks:
                raise ValueError("No chunks generated from document")
            
            # Add to ChromaDB
            chunk_count = add_to_chromadb(document_id, chunks)
            
            # Update document status
            db.documents.update_one(
                {"_id": ObjectId(document_id)},
                {
                    "$set": {
                        "status": DocumentStatus.PROCESSED.value,
                        "chunk_count": chunk_count,
                        "embedding_model": EMBEDDING_MODEL,
                        "processed_at": datetime.utcnow()
                    }
                }
            )
            
            logger.info(f"[OK] Document processed: {document_id} ({chunk_count} chunks)")
            return True
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
                
    except Exception as e:
        logger.error(f"Document processing failed: {str(e)}")
        db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {
                "$set": {
                    "status": DocumentStatus.FAILED.value,
                    "error_message": str(e)
                }
            }
        )
        return False


def extract_text(file_path: str, file_type: str) -> str:
    """
    Extract text content from a document
    
    Args:
        file_path: Path to the file
        file_type: File extension
    
    Returns:
        Extracted text content
    """
    text = ""
    
    try:
        if file_type == "pdf":
            # Use PyPDF2 or pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except ImportError:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        
        elif file_type == "docx":
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        elif file_type == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                
        elif file_type == "json":
            import json
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Flatten JSON to text
            text = json_to_text(data)
            
    except Exception as e:
        logger.error(f"Text extraction error: {str(e)}")
        raise
    
    return text.strip()


def json_to_text(data, prefix="") -> str:
    """Convert JSON data to readable text"""
    text = ""
    
    if isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, (dict, list)):
                text += json_to_text(value, f"{prefix}{key}: ")
            else:
                text += f"{prefix}{key}: {value}\n"
    elif isinstance(data, list):
        for i, item in enumerate(data):
            text += json_to_text(item, f"{prefix}[{i}] ")
    else:
        text += f"{prefix}{data}\n"
    
    return text


def split_into_chunks(text: str) -> list:
    """
    Split text into chunks for embedding
    
    Args:
        text: Full text content
    
    Returns:
        List of text chunks
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    
    logger.info(f"Split text into {len(chunks)} chunks")
    return chunks


def add_to_chromadb(document_id: str, chunks: list) -> int:
    """
    Add chunks to ChromaDB with embeddings
    
    Args:
        document_id: Document ID for metadata
        chunks: List of text chunks
    
    Returns:
        Number of chunks added
    """
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_chroma import Chroma
    
    # Initialize embeddings
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True}
    )
    
    # Initialize or load ChromaDB
    vectorstore = Chroma(
        collection_name="cui_knowledge_base",
        embedding_function=embeddings,
        persist_directory=CHROMA_PERSIST_DIRECTORY
    )
    
    # Add documents with metadata
    ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [{"document_id": document_id, "chunk_index": i} for i in range(len(chunks))]
    
    vectorstore.add_texts(
        texts=chunks,
        ids=ids,
        metadatas=metadatas
    )
    
    logger.info(f"[OK] Added {len(chunks)} chunks to ChromaDB for document {document_id}")
    
    return len(chunks)


# ===========================================
# Delete Document Embeddings
# ===========================================

def delete_document_embeddings(document_id: str) -> bool:
    """
    Delete all embeddings for a document from ChromaDB
    
    Args:
        document_id: Document ID
    
    Returns:
        True if successful
    """
    try:
        from langchain_huggingface import HuggingFaceEmbeddings
        from langchain_chroma import Chroma
        
        embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True}
        )
        
        vectorstore = Chroma(
            collection_name="cui_knowledge_base",
            embedding_function=embeddings,
            persist_directory=CHROMA_PERSIST_DIRECTORY
        )
        
        # Get all chunk IDs for this document
        # ChromaDB doesn't have direct filter delete, so we use where clause
        results = vectorstore._collection.get(
            where={"document_id": document_id}
        )
        
        if results and results["ids"]:
            vectorstore._collection.delete(ids=results["ids"])
            logger.info(f"[OK] Deleted {len(results['ids'])} embeddings for document {document_id}")
        
        return True
        
    except Exception as e:
        logger.error(f"Failed to delete embeddings: {str(e)}")
        return False


# ===========================================
# Rebuild All Embeddings
# ===========================================

def rebuild_all_embeddings() -> dict:
    """
    Rebuild all embeddings from uploaded documents
    
    Returns:
        Status report dictionary
    """
    db = get_database()
    
    results = {
        "processed": 0,
        "failed": 0,
        "errors": []
    }
    
    documents = db.documents.find()
    
    for doc in documents:
        try:
            # Get file from GridFS
            file_data = db.get_file(doc["gridfs_file_id"])
            if file_data:
                success = process_document(
                    str(doc["_id"]),
                    file_data,
                    doc["file_type"]
                )
                if success:
                    results["processed"] += 1
                else:
                    results["failed"] += 1
        except Exception as e:
            results["failed"] += 1
            results["errors"].append(f"{doc.get('file_name')}: {str(e)}")
    
    return results
