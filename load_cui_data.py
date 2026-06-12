"""
CUI Campus Bot Data Loader
Loads and processes CUI chatbot data from JSON files and PDF documents
"""

import json
import os
from typing import List, Dict, Optional
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    TokenTextSplitter
)


class TextChunker:
    """
    Advanced text chunking strategies for different types of content
    """
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None
    ):
        """
        Initialize text chunker with configurable parameters
        
        Args:
            chunk_size: Maximum size of each chunk
            chunk_overlap: Number of characters to overlap between chunks
            separators: Custom separators for splitting text
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Default separators optimized for Q&A and structured text
        if separators is None:
            self.separators = ["\n\n", "\n", ". ", " ", ""]
        else:
            self.separators = separators
        
        # Initialize different splitter types
        self.recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=self.separators
        )
        
        self.character_splitter = CharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator="\n\n"
        )
        
        self.token_splitter = TokenTextSplitter(
            chunk_size=chunk_size // 4,  # Approximate token count
            chunk_overlap=chunk_overlap // 4
        )
    
    def chunk_by_sentences(self, text: str, max_sentences: int = 5) -> List[str]:
        """
        Split text into chunks based on sentence boundaries
        
        Args:
            text: Text to split
            max_sentences: Maximum number of sentences per chunk
            
        Returns:
            List of text chunks
        """
        import re
        
        # Split by sentence boundaries
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = []
        
        for sentence in sentences:
            current_chunk.append(sentence)
            if len(current_chunk) >= max_sentences:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
        
        # Add remaining sentences
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def chunk_by_paragraphs(self, text: str) -> List[str]:
        """
        Split text into chunks based on paragraph boundaries
        
        Args:
            text: Text to split
            
        Returns:
            List of text chunks (paragraphs)
        """
        paragraphs = text.split('\n\n')
        chunks = [p.strip() for p in paragraphs if p.strip()]
        return chunks
    
    def chunk_with_context(
        self,
        text: str,
        metadata: Optional[Dict] = None
    ) -> List[Dict[str, str]]:
        """
        Create chunks while preserving context through overlap and metadata
        
        Args:
            text: Text to split
            metadata: Additional metadata to include
            
        Returns:
            List of dictionaries with chunk text and metadata
        """
        chunks = self.recursive_splitter.split_text(text)
        
        result = []
        for i, chunk in enumerate(chunks):
            chunk_data = {
                'text': chunk,
                'chunk_id': i,
                'total_chunks': len(chunks),
                'char_count': len(chunk)
            }
            
            if metadata:
                chunk_data.update(metadata)
            
            result.append(chunk_data)
        
        return result
    
    def smart_chunk_qna(self, question: str, answer: str) -> List[str]:
        """
        Intelligently chunk Q&A pairs
        - Keep short Q&A together
        - Split long answers while keeping question context
        
        Args:
            question: Question text
            answer: Answer text
            
        Returns:
            List of text chunks
        """
        full_text = f"Question: {question}\n\nAnswer: {answer}"
        
        # If total length is within chunk size, return as single chunk
        if len(full_text) <= self.chunk_size:
            return [full_text]
        
        # If answer is long, split it while keeping question prefix
        chunks = []
        answer_chunks = self.recursive_splitter.split_text(answer)
        
        for i, answer_chunk in enumerate(answer_chunks):
            if i == 0:
                # First chunk includes the full question
                chunk = f"Question: {question}\n\nAnswer: {answer_chunk}"
            else:
                # Subsequent chunks reference the question
                chunk = f"Question: {question} (continued)\n\nAnswer: {answer_chunk}"
            chunks.append(chunk)
        
        return chunks
    
    def chunk_documents(
        self,
        documents: List[Document],
        strategy: str = "recursive"
    ) -> List[Document]:
        """
        Chunk a list of Langchain Document objects
        
        Args:
            documents: List of Document objects to chunk
            strategy: Chunking strategy ('recursive', 'character', 'token')
            
        Returns:
            List of chunked Document objects
        """
        if strategy == "recursive":
            splitter = self.recursive_splitter
        elif strategy == "character":
            splitter = self.character_splitter
        elif strategy == "token":
            splitter = self.token_splitter
        else:
            raise ValueError(f"Unknown strategy: {strategy}")
        
        chunked_docs = splitter.split_documents(documents)
        
        # Add chunk information to metadata
        for i, doc in enumerate(chunked_docs):
            doc.metadata['chunk_id'] = i
            doc.metadata['char_count'] = len(doc.page_content)
        
        return chunked_docs


class CUIDataLoader:
    """
    Load and process CUI campus data from various sources
    """
    
    def __init__(
        self,
        data_directory: str = "cui_chatbot_data",
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        """
        Initialize the CUI data loader
        
        Args:
            data_directory: Path to directory containing CUI data files
            chunk_size: Maximum size of each text chunk
            chunk_overlap: Overlap between consecutive chunks
        """
        self.data_directory = data_directory
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text chunker
        self.chunker = TextChunker(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        # Keep the text splitter for PDF processing
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
    def load_json_qna(self, file_path: str, apply_chunking: bool = True) -> List[Document]:
        """
        Load Q&A data from JSON file and convert to Documents with optional chunking
        
        Args:
            file_path: Path to JSON file
            apply_chunking: Whether to apply smart chunking to Q&A pairs
            
        Returns:
            List of Document objects
        """
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            for item in data:
                category = item.get('category', 'General')
                question = item.get('question', '')
                answer = item.get('answer', '')
                
                metadata = {
                    'source': os.path.basename(file_path),
                    'category': category,
                    'type': 'qna',
                    'question': question
                }
                
                if apply_chunking:
                    # Use smart Q&A chunking
                    chunks = self.chunker.smart_chunk_qna(question, answer)
                    
                    for i, chunk_text in enumerate(chunks):
                        chunk_metadata = metadata.copy()
                        chunk_metadata['chunk_id'] = i
                        chunk_metadata['total_chunks'] = len(chunks)
                        
                        doc = Document(page_content=chunk_text, metadata=chunk_metadata)
                        documents.append(doc)
                else:
                    # Keep as single document
                    content = f"Question: {question}\n\nAnswer: {answer}"
                    doc = Document(page_content=content, metadata=metadata)
                    documents.append(doc)
            
            print(f"Loaded {len(documents)} document chunks from {os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"Error loading JSON file {file_path}: {str(e)}")
        
        return documents
    
    def load_pdf_timetable(self, file_path: str) -> List[Document]:
        """
        Load timetable PDF and split into chunks
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of Document objects
        """
        documents = []
        
        try:
            loader = PyPDFLoader(file_path)
            pdf_docs = loader.load()
            
            # Add metadata
            for doc in pdf_docs:
                doc.metadata['category'] = 'Timetable'
                doc.metadata['type'] = 'timetable'
                doc.metadata['source'] = os.path.basename(file_path)
            
            # Split into chunks
            documents = self.text_splitter.split_documents(pdf_docs)
            print(f"Loaded {len(documents)} chunks from {os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"Error loading PDF file {file_path}: {str(e)}")
        
        return documents
    
    def load_docx_file(self, file_path: str) -> List[Document]:
        """
        Load DOCX file and split into chunks
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of Document objects
        """
        documents = []
        
        try:
            from docx import Document as DocxDocument
            
            docx_doc = DocxDocument(file_path)
            full_text = []
            
            for para in docx_doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            content = "\n\n".join(full_text)
            
            # Create a Langchain Document
            doc = Document(
                page_content=content,
                metadata={
                    'source': os.path.basename(file_path),
                    'category': 'General',
                    'type': 'docx'
                }
            )
            
            # Split into chunks
            documents = self.text_splitter.split_documents([doc])
            
            # Update metadata for each chunk
            for i, chunk_doc in enumerate(documents):
                chunk_doc.metadata['chunk_id'] = i
            
            print(f"Loaded {len(documents)} chunks from {os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"Error loading DOCX file {file_path}: {str(e)}")
        
        return documents
    
    def load_all_data(self, mongo_db=None, mongo_gridfs=None) -> List[Document]:
        """
        Load all data from MongoDB GridFS or local directory
        
        Args:
            mongo_db: MongoDB database connection (optional)
            mongo_gridfs: MongoDB GridFS connection (optional)
        
        Returns:
            Combined list of all Document objects
        """
        all_documents = []
        
        print(f"\n{'='*60}")
        print(f"Loading CUI Campus Bot Data")
        print(f"{'='*60}\n")
        
        # Try to load from MongoDB GridFS first (PRIMARY source)
        if mongo_gridfs is not None and mongo_db is not None:
            print("📂 Loading documents from MongoDB Atlas GridFS...")
            mongodb_docs = self.load_from_mongodb(mongo_db, mongo_gridfs)
            if mongodb_docs:
                all_documents.extend(mongodb_docs)
                print(f"✅ Loaded {len(mongodb_docs)} chunks from MongoDB GridFS")
            else:
                print("⚠️ No documents found in MongoDB GridFS. Upload documents via Admin Dashboard!")
        else:
            # Fallback to local directory ONLY if MongoDB not configured
            print("⚠️ MongoDB not configured - falling back to local directory...")
            if os.path.exists(self.data_directory):
                local_docs = self._load_from_local_directory()
                if local_docs:
                    all_documents.extend(local_docs)
                    print(f"Loaded {len(local_docs)} chunks from local directory")
        
        if not all_documents:
            print("Warning: No documents found in MongoDB or local directory!")
        
        print(f"\n{'='*60}")
        print(f"Total documents loaded: {len(all_documents)}")
        print(f"{'='*60}\n")
        
        return all_documents
    
    def load_from_mongodb(self, mongo_db, mongo_gridfs) -> List[Document]:
        """
        Load documents from MongoDB GridFS
        
        Args:
            mongo_db: MongoDB database connection
            mongo_gridfs: MongoDB GridFS connection
        
        Returns:
            List of Document objects
        """
        import tempfile
        import io
        
        documents = []
        
        try:
            # Get all files from GridFS
            for grid_file in mongo_gridfs.find():
                filename = grid_file.filename
                file_content = grid_file.read()
                file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
                
                print(f"Processing from GridFS: {filename}")
                
                # Process based on file type
                if file_ext == 'pdf':
                    docs = self._process_pdf_content(file_content, filename)
                    documents.extend(docs)
                elif file_ext == 'docx':
                    docs = self._process_docx_content(file_content, filename)
                    documents.extend(docs)
                elif file_ext == 'txt':
                    docs = self._process_txt_content(file_content, filename)
                    documents.extend(docs)
                elif file_ext == 'json':
                    docs = self._process_json_content(file_content, filename)
                    documents.extend(docs)
                    
        except Exception as e:
            print(f"Error loading from MongoDB: {str(e)}")
        
        return documents
    
    def _process_pdf_content(self, content: bytes, filename: str) -> List[Document]:
        """Process PDF content from bytes"""
        import tempfile
        documents = []
        
        try:
            # Write to temp file for PyPDFLoader
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            # Load using PyPDFLoader
            loader = PyPDFLoader(tmp_path)
            pages = loader.load()
            
            for page in pages:
                page.metadata['source'] = filename
                page.metadata['category'] = 'General'
                page.metadata['type'] = 'pdf'
            
            documents = self.text_splitter.split_documents(pages)
            
            for i, doc in enumerate(documents):
                doc.metadata['chunk_id'] = i
            
            # Clean up temp file
            os.unlink(tmp_path)
            
            print(f"Loaded {len(documents)} chunks from {filename}")
            
        except Exception as e:
            print(f"Error processing PDF {filename}: {str(e)}")
        
        return documents
    
    def _process_docx_content(self, content: bytes, filename: str) -> List[Document]:
        """Process DOCX content from bytes"""
        import tempfile
        from docx import Document as DocxDocument
        import io
        
        documents = []
        
        try:
            # Load DOCX from bytes
            docx_file = DocxDocument(io.BytesIO(content))
            
            full_text = []
            for para in docx_file.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            text_content = "\n\n".join(full_text)
            
            doc = Document(
                page_content=text_content,
                metadata={
                    'source': filename,
                    'category': 'General',
                    'type': 'docx'
                }
            )
            
            documents = self.text_splitter.split_documents([doc])
            
            for i, chunk_doc in enumerate(documents):
                chunk_doc.metadata['chunk_id'] = i
            
            print(f"Loaded {len(documents)} chunks from {filename}")
            
        except Exception as e:
            print(f"Error processing DOCX {filename}: {str(e)}")
        
        return documents
    
    def _process_txt_content(self, content: bytes, filename: str) -> List[Document]:
        """Process TXT content from bytes"""
        documents = []
        
        try:
            text_content = content.decode('utf-8')
            
            doc = Document(
                page_content=text_content,
                metadata={
                    'source': filename,
                    'category': 'General',
                    'type': 'txt'
                }
            )
            
            documents = self.text_splitter.split_documents([doc])
            
            for i, chunk_doc in enumerate(documents):
                chunk_doc.metadata['chunk_id'] = i
            
            print(f"Loaded {len(documents)} chunks from {filename}")
            
        except Exception as e:
            print(f"Error processing TXT {filename}: {str(e)}")
        
        return documents
    
    def _process_json_content(self, content: bytes, filename: str) -> List[Document]:
        """Process JSON Q&A content from bytes"""
        import json
        documents = []
        
        try:
            data = json.loads(content.decode('utf-8'))
            
            for item in data:
                category = item.get('category', 'General')
                question = item.get('question', '')
                answer = item.get('answer', '')
                
                content_text = f"Question: {question}\n\nAnswer: {answer}"
                
                doc = Document(
                    page_content=content_text,
                    metadata={
                        'source': filename,
                        'category': category,
                        'type': 'qna',
                        'question': question
                    }
                )
                documents.append(doc)
            
            print(f"Loaded {len(documents)} Q&A items from {filename}")
            
        except Exception as e:
            print(f"Error processing JSON {filename}: {str(e)}")
        
        return documents
    
    def _load_from_local_directory(self) -> List[Document]:
        """Load documents from local directory (backward compatibility)"""
        all_documents = []
        
        # Load all JSON files dynamically
        json_files = [f for f in os.listdir(self.data_directory) if f.lower().endswith('.json')]
        if not json_files:
            print("Warning: No JSON files found in data directory")
        for json_file in json_files:
            file_path = os.path.join(self.data_directory, json_file)
            docs = self.load_json_qna(file_path)
            all_documents.extend(docs)

        # Load all PDFs dynamically
        pdf_files = [f for f in os.listdir(self.data_directory) if f.lower().endswith('.pdf')]
        if not pdf_files:
            print("Warning: No PDF files found in data directory")
        for pdf_file in pdf_files:
            pdf_path = os.path.join(self.data_directory, pdf_file)
            docs = self.load_pdf_timetable(pdf_path)
            all_documents.extend(docs)
        
        # Load all DOCX files dynamically
        docx_files = [f for f in os.listdir(self.data_directory) if f.lower().endswith('.docx')]
        if docx_files:
            print(f"Found {len(docx_files)} DOCX files")
        for docx_file in docx_files:
            docx_path = os.path.join(self.data_directory, docx_file)
            docs = self.load_docx_file(docx_path)
            all_documents.extend(docs)
        
        print(f"\n{'='*60}")
        print(f"Total documents loaded: {len(all_documents)}")
        print(f"{'='*60}\n")
        
        return all_documents
    
    def get_data_summary(self) -> Dict[str, int]:
        """
        Get summary statistics of loaded data
        
        Returns:
            Dictionary with category counts
        """
        documents = self.load_all_data()
        
        summary = {}
        for doc in documents:
            category = doc.metadata.get('category', 'Unknown')
            summary[category] = summary.get(category, 0) + 1
        
        return summary
    
    def get_chunking_stats(self, documents: List[Document]) -> Dict:
        """
        Get detailed statistics about text chunking
        
        Args:
            documents: List of documents to analyze
            
        Returns:
            Dictionary with chunking statistics
        """
        stats = {
            'total_documents': len(documents),
            'total_characters': 0,
            'avg_chunk_size': 0,
            'min_chunk_size': float('inf'),
            'max_chunk_size': 0,
            'chunks_by_category': {}
        }
        
        for doc in documents:
            char_count = len(doc.page_content)
            stats['total_characters'] += char_count
            stats['min_chunk_size'] = min(stats['min_chunk_size'], char_count)
            stats['max_chunk_size'] = max(stats['max_chunk_size'], char_count)
            
            category = doc.metadata.get('category', 'Unknown')
            stats['chunks_by_category'][category] = stats['chunks_by_category'].get(category, 0) + 1
        
        if documents:
            stats['avg_chunk_size'] = stats['total_characters'] / len(documents)
        
        return stats


# Example usage and testing
if __name__ == "__main__":
    # Initialize the loader with custom chunk size
    loader = CUIDataLoader(
        data_directory="cui_chatbot_data",
        chunk_size=800,
        chunk_overlap=150
    )
    
    # Load all data
    documents = loader.load_all_data()
    
    # Display summary
    print("\nData Summary by Category:")
    print("-" * 40)
    summary = {}
    for doc in documents:
        category = doc.metadata.get('category', 'Unknown')
        summary[category] = summary.get(category, 0) + 1
    
    for category, count in sorted(summary.items()):
        print(f"{category}: {count} documents")
    
    # Display chunking statistics
    print("\n" + "="*60)
    print("Chunking Statistics:")
    print("="*60)
    stats = loader.get_chunking_stats(documents)
    print(f"Total documents/chunks: {stats['total_documents']}")
    print(f"Total characters: {stats['total_characters']:,}")
    print(f"Average chunk size: {stats['avg_chunk_size']:.0f} characters")
    print(f"Min chunk size: {stats['min_chunk_size']} characters")
    print(f"Max chunk size: {stats['max_chunk_size']} characters")
    
    print("\nChunks by category:")
    for category, count in sorted(stats['chunks_by_category'].items()):
        print(f"  {category}: {count} chunks")
    
    # Display sample documents
    print("\n" + "="*60)
    print("Sample Document Chunks:")
    print("="*60)
    
    for i, doc in enumerate(documents[:3]):
        print(f"\nChunk {i+1}:")
        print(f"Category: {doc.metadata.get('category')}")
        print(f"Source: {doc.metadata.get('source')}")
        print(f"Chunk ID: {doc.metadata.get('chunk_id', 'N/A')}")
        print(f"Character count: {len(doc.page_content)}")
        print(f"Content preview: {doc.page_content[:250]}...")
        print("-" * 60)
    
    # Test different chunking strategies
    print("\n" + "="*60)
    print("Testing Different Chunking Strategies:")
    print("="*60)
    
    sample_text = """
    COMSATS University Islamabad is a premier institution of higher learning.
    The university offers various undergraduate and graduate programs.
    It has multiple campuses across Pakistan including Islamabad, Lahore, and Abbottabad.
    The admission process is merit-based and conducted twice a year.
    """
    
    chunker = loader.chunker
    
    print("\n1. Sentence-based chunking (max 2 sentences per chunk):")
    sentence_chunks = chunker.chunk_by_sentences(sample_text.strip(), max_sentences=2)
    for i, chunk in enumerate(sentence_chunks):
        print(f"   Chunk {i+1}: {chunk}")
    
    print("\n2. Paragraph-based chunking:")
    para_chunks = chunker.chunk_by_paragraphs(sample_text)
    print(f"   Number of paragraphs: {len(para_chunks)}")
    
    print("\n3. Context-aware chunking:")
    context_chunks = chunker.chunk_with_context(
        sample_text.strip(),
        metadata={'source': 'test', 'category': 'General'}
    )
    for chunk_data in context_chunks:
        print(f"   Chunk {chunk_data['chunk_id']+1}/{chunk_data['total_chunks']}: "
              f"{chunk_data['char_count']} chars")
